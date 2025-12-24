# table_aware_loader.py

"""
Table-Aware Document Loader

Enhanced PDF loader that specifically handles tabular data (like fee structures)
by preserving row-column relationships and creating structured text representations.

This solves the problem where fee tables lose their structure when extracted as plain text.
"""

import os
import re
from typing import Dict, List, Tuple, Optional
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

from logging_config import get_logger

logger = get_logger(__name__)


class TableAwareLoader:
    """
    Document loader that preserves table structure for better RAG retrieval.
    
    Key improvements:
    1. Detects and extracts tables with structure preservation
    2. Converts tables to structured text with clear row-column relationships
    3. Creates semantic chunks that keep related data together (e.g., program + fee)
    """
    
    def __init__(self, dpi: float = 2.0):
        self.dpi = dpi
        self.tesseract_config = r'--oem 3 --psm 6'
    
    def extract_tables_from_page(self, page: fitz.Page) -> List[Dict]:
        """
        Extract tables from a PDF page.
        
        IMPORTANT: For tuition fee documents, we prioritize text-based extraction
        because PyMuPDF often extracts these tables poorly (split headers, wrong structure).
        
        Returns:
            List of table dictionaries with headers and rows
        """
        # First, try text-based extraction which works better for fee tables
        text_tables = self._extract_tables_from_text(page)
        
        # If we found fee tables with actual programme data, use those
        if text_tables:
            for table in text_tables:
                if table.get('rows') and len(table['rows']) > 0:
                    # Check if this looks like a fee table (has programme names)
                    first_row = table['rows'][0]
                    if first_row and any('Bachelor' in str(v) or 'Master' in str(v) or 'Foundation' in str(v) 
                                        for v in first_row if v):
                        return text_tables
        
        # Fallback to PyMuPDF table detection
        tables = []
        try:
            page_tables = page.find_tables()
            table_list = page_tables.tables if hasattr(page_tables, 'tables') else list(page_tables)
            
            for table in table_list:
                extracted = table.extract()
                if extracted and len(extracted) > 0:
                    headers = extracted[0] if extracted else []
                    rows = extracted[1:] if len(extracted) > 1 else []
                    
                    tables.append({
                        'headers': headers,
                        'rows': rows,
                        'bbox': table.bbox if hasattr(table, 'bbox') else None
                    })
        except Exception as e:
            logger.debug(f"PyMuPDF table extraction failed: {e}")
        
        # If PyMuPDF found tables, use them; otherwise use text-based
        return tables if tables else text_tables
    
    def _extract_tables_from_text(self, page: fitz.Page) -> List[Dict]:
        """
        Enhanced fallback table extraction from text layout.
        Specifically designed for tuition fee tables and similar structured data.
        """
        tables = []
        text = page.get_text("text")
        lines = text.split('\n')
        
        # Strategy 1: Look for fee table patterns (Programme + Duration + Fee)
        # This is the primary strategy for tuition fee documents
        fee_table = self._extract_fee_table_pattern(lines)
        if fee_table and fee_table.get('rows'):
            tables.append(fee_table)
            # If we found a fee table, skip generic extraction to avoid duplicates
            return tables
        
        # Strategy 2: Generic multi-column detection (only if no fee table found)
        generic_tables = self._extract_generic_tables(lines)
        tables.extend(generic_tables)
        
        return tables
    
    def _extract_fee_table_pattern(self, lines: List[str]) -> Optional[Dict]:
        """
        Extract fee tables with pattern: Programme name followed by duration and fee.
        Handles the specific format in XMUM tuition fee PDFs.
        """
        rows = []
        headers = ['Programme', 'Duration', 'Tuition Fee (RM per year)']
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Look for programme names (Bachelor, Master, Foundation, etc.)
            if re.match(r'^(Bachelor|Master|Doctor|Foundation|Diploma|PhD)\s', line, re.IGNORECASE):
                programme_name = line
                
                # Skip registration codes like [R2/0411/6/0017]
                j = i + 1
                while j < len(lines) and re.match(r'^\[.*\]', lines[j].strip()):
                    j += 1
                
                # Look for duration and fee in next few lines
                duration = None
                fee = None
                
                for k in range(j, min(j + 5, len(lines))):
                    next_line = lines[k].strip()
                    
                    # Match duration patterns like "4 years", "1 year", "2.5 years"
                    duration_match = re.search(r'(\d+(?:\.\d+)?\s*years?)', next_line, re.IGNORECASE)
                    if duration_match and not duration:
                        duration = duration_match.group(1)
                    
                    # Match fee patterns like "31,000" or "19,000"
                    fee_match = re.search(r'^([\d,]+)$', next_line)
                    if fee_match and not fee:
                        fee = fee_match.group(1)
                    
                    # Also check for combined duration-fee line
                    combined_match = re.match(r'^(\d+(?:\.\d+)?\s*years?)\s+([\d,]+)', next_line, re.IGNORECASE)
                    if combined_match:
                        duration = combined_match.group(1)
                        fee = combined_match.group(2)
                
                if programme_name and (duration or fee):
                    rows.append([programme_name, duration or '-', fee or '-'])
                
                i = j
            else:
                i += 1
        
        if rows:
            return {
                'headers': headers,
                'rows': rows,
                'bbox': None
            }
        return None
    
    def _extract_generic_tables(self, lines: List[str]) -> List[Dict]:
        """
        Extract generic tables based on column alignment.
        """
        tables = []
        potential_table_rows = []
        
        for line in lines:
            # Split by multiple spaces or tabs
            parts = re.split(r'\s{2,}|\t+', line.strip())
            if len(parts) >= 2:  # At least 2 columns
                potential_table_rows.append(parts)
        
        if len(potential_table_rows) >= 2:
            # Group consecutive rows with similar column counts
            current_table = []
            for row in potential_table_rows:
                if not current_table or abs(len(row) - len(current_table[0])) <= 1:
                    current_table.append(row)
                else:
                    if len(current_table) >= 2:
                        tables.append({
                            'headers': current_table[0],
                            'rows': current_table[1:],
                            'bbox': None
                        })
                    current_table = [row]
            
            if len(current_table) >= 2:
                tables.append({
                    'headers': current_table[0],
                    'rows': current_table[1:],
                    'bbox': None
                })
        
        return tables
    
    def table_to_structured_text(self, table: Dict, context: str = "") -> str:
        """
        Convert a table to structured, searchable text.
        
        Each row becomes a complete, self-contained statement that includes
        the column headers for context.
        
        Example:
            Headers: [Programme, Duration, Fee]
            Row: [Bachelor of Data Science, 4 years, RM 31,000]
            
            Output: "Programme: Bachelor of Data Science | Duration: 4 years | Annual Fee: RM 31,000"
        """
        if not table.get('headers') or not table.get('rows'):
            return ""
        
        headers = [str(h).strip() for h in table['headers'] if h]
        structured_lines = []
        
        # Add context header if provided
        if context:
            structured_lines.append(f"[{context}]")
        
        for row in table['rows']:
            row_values = [str(v).strip() for v in row if v]
            
            # Skip empty rows
            if not any(row_values):
                continue
            
            # Create structured representation
            parts = []
            for i, value in enumerate(row_values):
                if i < len(headers) and headers[i]:
                    parts.append(f"{headers[i]}: {value}")
                else:
                    parts.append(value)
            
            if parts:
                structured_lines.append(" | ".join(parts))
        
        return "\n".join(structured_lines)
    
    def create_semantic_chunks_from_table(
        self, 
        table: Dict, 
        source_file: str,
        page_num: int,
        context: str = ""
    ) -> List[Dict]:
        """
        Create semantic chunks from a table where each chunk contains
        complete information about a single item (e.g., one program's fees).
        
        This ensures that when searching for "Bachelor of Data Science fees",
        the complete fee information is in a single retrievable chunk.
        """
        chunks = []
        
        if not table.get('headers') or not table.get('rows'):
            return chunks
        
        headers = [str(h).strip() for h in table['headers'] if h]
        
        for row_idx, row in enumerate(table['rows']):
            row_values = [str(v).strip() for v in row if v]
            
            # Skip empty rows
            if not any(row_values):
                continue
            
            # Create a complete, self-contained chunk for this row
            chunk_parts = []
            
            # Add context
            if context:
                chunk_parts.append(f"Document: {context}")
            
            chunk_parts.append(f"Source: {source_file}, Page {page_num}")
            
            # Add all column-value pairs
            for i, value in enumerate(row_values):
                if i < len(headers) and headers[i]:
                    chunk_parts.append(f"{headers[i]}: {value}")
                else:
                    chunk_parts.append(f"Value: {value}")
            
            chunk_text = "\n".join(chunk_parts)
            
            chunks.append({
                'text': chunk_text,
                'metadata': {
                    'source': source_file,
                    'page': page_num,
                    'row_index': row_idx,
                    'type': 'table_row',
                    'headers': headers
                }
            })
        
        return chunks
    
    def load_pdf_with_tables(self, filepath: str) -> Tuple[str, List[Dict]]:
        """
        Load a PDF with enhanced table extraction.
        
        Returns:
            Tuple of (regular_text, list_of_table_chunks)
        """
        doc = fitz.open(filepath)
        filename = os.path.basename(filepath)
        
        logger.info(f"Processing PDF with table awareness: {filename}")
        
        all_text_parts = []
        all_table_chunks = []
        
        for page_num, page in enumerate(doc, 1):
            # Extract tables first
            tables = self.extract_tables_from_page(page)
            
            # Create semantic chunks from tables
            for table in tables:
                # Determine context from page content
                page_text = page.get_text()
                context = self._extract_context(page_text, filename)
                
                chunks = self.create_semantic_chunks_from_table(
                    table, filename, page_num, context
                )
                all_table_chunks.extend(chunks)
                
                # Also add structured table text
                structured_text = self.table_to_structured_text(table, context)
                if structured_text:
                    all_text_parts.append(f"\n[Table from Page {page_num}]\n{structured_text}")
            
            # Extract regular text (excluding table areas if possible)
            regular_text = page.get_text()
            if regular_text.strip():
                all_text_parts.append(f"\n[Page {page_num}]\n{regular_text}")
        
        doc.close()
        
        full_text = "\n".join(all_text_parts)
        
        logger.info(
            f"PDF processed: {filename}",
            extra={
                'pages': page_num,
                'table_chunks': len(all_table_chunks),
                'text_length': len(full_text)
            }
        )
        
        return full_text, all_table_chunks
    
    def _extract_context(self, page_text: str, filename: str) -> str:
        """
        Extract context (e.g., 'Local Tuition Fees 2026') from page text or filename.
        """
        # Try to find a title-like pattern at the start
        lines = page_text.strip().split('\n')[:5]  # First 5 lines
        
        for line in lines:
            line = line.strip()
            # Look for title patterns
            if re.match(r'^[A-Z][A-Za-z\s]+(?:Fees?|Tuition|Scholarship|Programme)', line):
                return line[:100]  # Limit length
        
        # Fall back to filename
        name = os.path.splitext(filename)[0]
        return name.replace('_', ' ').replace('-', ' ')
    
    def load_folder_with_tables(self, folder_path: str) -> Tuple[Dict[str, str], List[Dict]]:
        """
        Load all documents from a folder with table awareness.
        
        Returns:
            Tuple of (documents_dict, all_table_chunks)
        """
        documents = {}
        all_chunks = []
        
        for filename in os.listdir(folder_path):
            filepath = os.path.join(folder_path, filename)
            
            if not os.path.isfile(filepath):
                continue
            
            ext = os.path.splitext(filename)[1].lower()
            
            if ext == '.pdf':
                text, table_chunks = self.load_pdf_with_tables(filepath)
                documents[filename] = text
                all_chunks.extend(table_chunks)
            elif ext == '.txt':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    documents[filename] = f.read()
            elif ext == '.docx':
                try:
                    import docx
                    doc = docx.Document(filepath)
                    text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    documents[filename] = text
                except ImportError:
                    logger.warning(f"python-docx not installed, skipping {filename}")
        
        logger.info(
            f"Folder loaded with table awareness",
            extra={
                'documents': len(documents),
                'table_chunks': len(all_chunks)
            }
        )
        
        return documents, all_chunks


def extract_fee_structure(text: str) -> List[Dict]:
    """
    Specialized extractor for fee structures.
    
    Looks for patterns like:
    - "RM XX,XXX" or "RM XX,XXX.XX"
    - Programme names followed by fees
    - Duration information (e.g., "4 years")
    """
    fee_entries = []
    
    # Pattern for Malaysian Ringgit amounts
    fee_pattern = r'RM\s*[\d,]+(?:\.\d{2})?'
    
    # Pattern for programme names (typically capitalized words)
    programme_pattern = r'(?:Bachelor|Master|Doctor|PhD|Diploma|Certificate)\s+(?:of|in)\s+[A-Za-z\s\(\)]+(?:Honours?)?'
    
    # Pattern for duration
    duration_pattern = r'(\d+)\s*(?:years?|semesters?)'
    
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        # Look for fee amounts
        fee_match = re.search(fee_pattern, line)
        if fee_match:
            fee = fee_match.group()
            
            # Look for programme name in nearby lines
            context_start = max(0, i - 3)
            context_end = min(len(lines), i + 2)
            context = ' '.join(lines[context_start:context_end])
            
            programme_match = re.search(programme_pattern, context, re.IGNORECASE)
            duration_match = re.search(duration_pattern, context, re.IGNORECASE)
            
            entry = {
                'fee': fee,
                'programme': programme_match.group() if programme_match else None,
                'duration': duration_match.group() if duration_match else None,
                'context': context[:500]
            }
            
            fee_entries.append(entry)
    
    return fee_entries


if __name__ == "__main__":
    import sys
    from logging_config import setup_logging
    
    setup_logging(level="INFO", json_output=False, app_name="table-loader")
    
    loader = TableAwareLoader()
    
    if len(sys.argv) > 1:
        path = sys.argv[1]
        if os.path.isfile(path) and path.endswith('.pdf'):
            text, chunks = loader.load_pdf_with_tables(path)
            print(f"\nExtracted {len(chunks)} table chunks")
            for chunk in chunks[:5]:
                print(f"\n--- Chunk ---\n{chunk['text']}")
        elif os.path.isdir(path):
            docs, chunks = loader.load_folder_with_tables(path)
            print(f"\nLoaded {len(docs)} documents with {len(chunks)} table chunks")
    else:
        print("Usage: python table_aware_loader.py <pdf_file_or_folder>")
